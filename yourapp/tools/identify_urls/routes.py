import os, re, io, shutil, tempfile, traceback, gc
from pathlib import Path
from typing import List, Set

from flask import Blueprint, render_template, request, send_file, abort
from werkzeug.utils import secure_filename

# PDF text extraction
from pypdf import PdfReader

# DOCX writing
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.opc.constants import RELATIONSHIP_TYPE as RT

# Safe escaping for ReportLab markup
from xml.sax.saxutils import escape as xml_escape

bp = Blueprint("identify_urls", __name__, url_prefix="/identify-urls")

# ------------------ Helpers ------------------

def is_pdf(name: str) -> bool:
    return name.lower().endswith(".pdf")

# Pragmatic URL matcher (http/https/www), trims trailing punctuation.
URL_PATTERN = re.compile(
    r"""(?ix)
    \b(
      (?:https?://|http://|www\.)       # scheme or www
      [^\s<>()\[\]{}"']+               # body
      (?:/[^\s<>()\[\]{}"']*)?         # optional path/query
    )
    """
)
TRIM_TRAILING = ".,);:!?'\"”’"

def extract_urls_from_text(text: str) -> List[str]:
    if not text:
        return []
    found = []
    for m in URL_PATTERN.finditer(text):
        u = m.group(1).strip()
        while u and u[-1] in TRIM_TRAILING:
            u = u[:-1]
        found.append(u)
    return found

def normalize_url(u: str) -> str:
    u = (u or "").strip()
    if u.lower().startswith(("http://", "https://")):
        return u
    if u.lower().startswith("www."):
        return "https://" + u
    return "https://" + u

def add_hyperlink(paragraph, url: str, text: str = None):
    """
    Preferred: real external hyperlink run using a relationship (blue, underlined in Word).
    """
    url = normalize_url(url)
    display = text or url

    part = paragraph.part
    r_id = part.relate_to(url, RT.HYPERLINK, is_external=True)

    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)

    r = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    rStyle = OxmlElement("w:rStyle")
    rStyle.set(qn("w:val"), "Hyperlink")
    rPr.append(rStyle)
    r.append(rPr)

    t = OxmlElement("w:t")
    t.set(qn("xml:space"), "preserve")
    t.text = display
    r.append(t)

    hyperlink.append(r)
    paragraph._p.append(hyperlink)
    return paragraph

def add_hyperlink_field(paragraph, url: str, text: str = None):
    """
    Fallback: field code HYPERLINK (very compatible with Word/LibreOffice).
    """
    url = normalize_url(url)
    display = text or url

    def _r(txt=None, instr=False, sep=False, fldCharType=None):
        r = OxmlElement("w:r")
        if instr:
            rPr = OxmlElement("w:rPr")
            rStyle = OxmlElement("w:rStyle")
            rStyle.set(qn("w:val"), "Hyperlink")
            rPr.append(rStyle)
            r.append(rPr)
            instrText = OxmlElement("w:instrText")
            instrText.set(qn("xml:space"), "preserve")
            instrText.text = txt or ""
            r.append(instrText)
            return r
        if sep:
            fldChar = OxmlElement("w:fldChar")
            fldChar.set(qn("w:fldCharType"), "separate")
            r.append(fldChar)
            return r
        if fldCharType:
            fldChar = OxmlElement("w:fldChar")
            fldChar.set(qn("w:fldCharType"), fldCharType)  # "begin" or "end"
            r.append(fldChar)
            return r
        t = OxmlElement("w:t")
        t.set(qn("xml:space"), "preserve")
        t.text = txt or ""
        r.append(t)
        return r

    p = paragraph._p
    p.append(_r(fldCharType="begin"))
    p.append(_r(txt=f'HYPERLINK "{url}"', instr=True))
    p.append(_r(sep=True))
    p.append(_r(display))
    p.append(_r(fldCharType="end"))
    return paragraph

def read_pdf_text_urls(pdf_path: str) -> Set[str]:
    """
    Safely open PDF, handle encryption if possible, extract text and URLs.
    """
    urls: Set[str] = set()
    try:
        reader = PdfReader(pdf_path, strict=False)
    except Exception as ex:
        print(f"[IDENTIFY_URLS] PdfReader open error: {ex}")
        raise

    # Try to open encrypted PDFs with empty password
    if getattr(reader, "is_encrypted", False):
        try:
            if reader.decrypt("") == 0:
                # can't decrypt—stop with a 400 so user gets a helpful message
                abort(400, "This PDF is encrypted. Please upload an unprotected file.")
            else:
                print("[IDENTIFY_URLS] decrypted with empty password")
        except Exception as ex:
            print(f"[IDENTIFY_URLS] decrypt error: {ex}")
            abort(400, "This PDF is encrypted and could not be opened.")

    total_pages = len(reader.pages)
    print(f"[IDENTIFY_URLS] total_pages={total_pages}")
    if total_pages < 1:
        abort(400, "The uploaded PDF appears to be empty.")

    for i in range(total_pages):
        try:
            text = (reader.pages[i].extract_text() or "")
            for u in extract_urls_from_text(text):
                urls.add(u)
        except Exception as ex:
            print(f"[IDENTIFY_URLS] page {i+1} extract error: {ex}")
            continue

    return urls

# ------------------ Routes ------------------

@bp.route("", methods=["GET"])
@bp.route("/", methods=["GET"])
def urls_get():
    return render_template("identify_urls.html")

@bp.route("", methods=["POST"])
@bp.route("/", methods=["POST"])
def urls_post():
    try:
        f = request.files.get("file")
        if not f or not f.filename:
            abort(400, "Please upload a PDF file.")
        name = secure_filename(f.filename)
        if not is_pdf(name):
            abort(400, "Only PDF files are accepted.")

        workdir = tempfile.mkdtemp(prefix="urls_")
        try:
            pdf_path = os.path.join(workdir, name)
            f.save(pdf_path)

            urls = read_pdf_text_urls(pdf_path)
            stem = Path(name).stem
            ofmt = (request.form.get("output_format") or "pdf").lower()
            print(f"[IDENTIFY_URLS] urls_found={len(urls)} ofmt={ofmt}")

            # ---- PDF OUTPUT (default; clickable links via ReportLab) ----
            if ofmt == "pdf":
                try:
                    # Lazy-import to reduce baseline RAM and avoid ImportError affecting DOCX path
                    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
                    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
                    from reportlab.lib.pagesizes import A4
                    from reportlab.lib.units import mm
                    from reportlab.lib.colors import blue

                    buf = io.BytesIO()
                    doc = SimpleDocTemplate(
                        buf, pagesize=A4,
                        leftMargin=20*mm, rightMargin=20*mm, topMargin=20*mm, bottomMargin=20*mm
                    )
                    styles = getSampleStyleSheet()
                    h1 = styles["Heading1"]
                    body = styles["BodyText"]
                    link_style = ParagraphStyle("LinkBody", parent=body, textColor=blue)

                    story = []
                    story.append(Paragraph(f'URLs found in "{stem}.pdf"', h1))
                    if urls:
                        story.append(Paragraph(f"Total unique URLs: {len(urls)}", body))
                        story.append(Spacer(1, 6))
                        for u in sorted(urls, key=str.lower):
                            nu = normalize_url(u)
                            href = xml_escape(nu, {'"': '&quot;'})
                            text = xml_escape(nu)
                            try:
                                story.append(Paragraph(f'<link href="{href}">{text}</link>', link_style))
                            except Exception as ex:
                                print(f"[IDENTIFY_URLS link escape fallback] {nu}: {ex}")
                                story.append(Paragraph(text, body))
                            story.append(Spacer(1, 2))
                    else:
                        story.append(Paragraph("No URLs were found in the text of this PDF.", body))
                        story.append(Paragraph("(Note: scanned PDFs/images need OCR; text-only extraction was used.)", body))

                    doc.build(story)
                    buf.seek(0)
                    return send_file(
                        buf,
                        as_attachment=True,
                        download_name=f"{stem}_URLs.pdf",
                        mimetype="application/pdf",
                    )
                except ImportError as ie:
                    print(f"[IDENTIFY_URLS] ReportLab not available: {ie}")
                    # if ReportLab missing, silently fall back to DOCX below

            # ---- DOCX OUTPUT (clickable links; robust fallback) ----
            docx_buf = io.BytesIO()
            doc = Document()
            doc.add_heading(f'URLs found in "{stem}.pdf"', level=1)
            if urls:
                doc.add_paragraph(f"Total unique URLs: {len(urls)}")
                doc.add_paragraph(
                    "(Links below are inserted as clickable hyperlinks. "
                    "If your viewer disables them, they will still display as plain URLs.)"
                )
                for u in sorted(urls, key=str.lower):
                    p = doc.add_paragraph()
                    try:
                        add_hyperlink(p, u)             # primary
                    except Exception as ex:
                        print(f"[IDENTIFY_URLS primary failed] {u}: {ex}")
                        try:
                            add_hyperlink_field(p, u)    # fallback
                        except Exception as ex2:
                            print(f"[IDENTIFY_URLS fallback failed] {u}: {ex2}")
                            p.add_run(normalize_url(u))  # last resort
            else:
                doc.add_paragraph("No URLs were found in the text of this PDF.")
                doc.add_paragraph("(Note: scanned PDFs/images need OCR; text-only extraction was used.)")

            out_name = f"{stem}_URLs.docx"
            doc.save(docx_buf)
            docx_buf.seek(0)
            return send_file(
                docx_buf,
                as_attachment=True,
                download_name=out_name,
                mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )

        finally:
            shutil.rmtree(workdir, ignore_errors=True)
            gc.collect()

    except Exception as e:
        print("[/identify-urls ERROR]", repr(e))
        traceback.print_exc()
        abort(500, "An error occurred while identifying URLs. Please try another file.")
